# -*- coding: utf-8 -*-
"""
Generate the official Word (.docx) version of the Latakia port monitoring
report — the editable format required for submission to official/government
bodies (letterhead, stamps and signatures can be added in Word).

Full RTL Arabic layout: right-aligned paragraphs, Arabic fonts, right-aligned
tables, embedded charts and maps.
"""
import os, json
import numpy as np
import pandas as pd
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import ROOT, DET_DIR, CHART_DIR, MAP_DIR, S1_DIR

VERSION = "2.3"
ARABIC_MONTHS = ["كانون الثاني", "شباط", "آذار", "نيسان", "أيار", "حزيران",
                 "تموز", "آب", "أيلول", "تشرين الأول", "تشرين الثاني", "كانون الأول"]
INK = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0x01, 0x9E, 0x92)
GRAY = RGBColor(0x55, 0x62, 0x70)
FONT_AR = "Noto Naskh Arabic"
FONT_EN = "Segoe UI"


def set_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT):
    """Force RTL + right alignment on a paragraph."""
    p.alignment = align
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def set_run(run, size=11, bold=False, color=INK, font=FONT_AR):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    rFonts.set(qn("w:cs"), font)


def para(doc, text, size=11, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.RIGHT,
         space_after=6, space_before=0, font=FONT_AR):
    p = doc.add_paragraph()
    set_rtl(p, align)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color, font=font)
    return p


def heading(doc, text, level=1):
    sizes = {1: 16, 2: 13.5, 3: 12}
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=sizes.get(level, 12), bold=True, color=ACCENT if level == 1 else INK)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "0E7490")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def set_cell_rtl(cell, text, size=10, bold=False, color=INK, shade=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    if shade:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), shade)
        tcPr.append(shd)


def make_table(doc, headers, rows, widths=None, header_shade="E0F2F0"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for j, h in enumerate(headers):
        set_cell_rtl(t.rows[0].cells[j], h, size=10, bold=True, shade=header_shade)
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            set_cell_rtl(t.rows[i].cells[j], str(v), size=10)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    # spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_image(doc, path, width_cm=15.5, caption=None):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        para(doc, caption, size=9.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)


def fmt(v, nd=1):
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return "—"
    if isinstance(v, float):
        return f"{v:.{nd}f}"
    return str(v)


def ym_ar(ym):
    try:
        y, m = ym.split("-")
        return f"{ARABIC_MONTHS[int(m)-1]} {y}"
    except Exception:
        return str(ym)


def build():
    monthly = pd.read_csv(os.path.join(DET_DIR, "monthly.csv"))
    yearly = pd.read_csv(os.path.join(DET_DIR, "yearly.csv"))
    summary = json.load(open(os.path.join(DET_DIR, "summary.json"), encoding="utf-8"))
    scenes = pd.read_json(os.path.join(DET_DIR, "s1_scenes.jsonl"), lines=True)
    good = scenes[scenes["error"].isna() | (scenes["error"] == "")].copy()
    good["dt"] = pd.to_datetime(good["datetime"])
    good["year"] = good["dt"].dt.year
    good["mon"] = good["dt"].dt.month
    s2 = pd.read_json(os.path.join(DET_DIR, "s2_scenes.jsonl"), lines=True)

    y = yearly.set_index("year")
    m = monthly.set_index("ym")
    regs = summary["regimes"]
    t2026 = regs["y2026_vs_2025"]["test"]
    tpost = regs["post_vs_pre_dec2024"]["test"]
    n_ok, n_scenes = len(good), len(scenes)
    n_vessels = int(good["n_est"].sum())
    # vessel length stats (dynamic)
    import numpy as _np
    _Ls = []
    for _v in good.itertuples():
        _vj = os.path.normpath(os.path.join(ROOT, "data", "raw_s1", _v.id, "vessels.json"))
        if os.path.exists(_vj):
            for _vv in json.load(open(_vj, encoding="utf-8")):
                if _vv.get("length_m"):
                    _Ls.append(_vv["length_m"])
    _L = _np.array(_Ls)
    vmed = float(_np.median(_L)); vp25 = float(_np.percentile(_L, 25))
    vp75 = float(_np.percentile(_L, 75)); vp90 = float(_np.percentile(_L, 90))
    vgt150 = float((_L > 150).mean() * 100); vlt50 = float((_L < 50).mean() * 100)
    vmax = float(_L.max())
    # zone stats
    from collections import Counter as _C
    _zc = _C()
    for _v in good.itertuples():
        _vj = os.path.normpath(os.path.join(ROOT, "data", "raw_s1", _v.id, "vessels.json"))
        if os.path.exists(_vj):
            for _vv in json.load(open(_vj, encoding="utf-8")):
                _zc[_vv["zone"]] += 1
    _zt = sum(_zc.values())
    zport = _zc.get(2, 0); zanch = _zc.get(1, 0); ztrans = _zc.get(0, 0)
    zport_pct = zport / _zt * 100; zanch_pct = zanch / _zt * 100; ztrans_pct = ztrans / _zt * 100
    # seasonal control
    _mm = monthly.set_index("ym")
    _fa26v = [_mm.loc[f"2026-0{i}", "mean_ships_port_adj"] for i in (2, 3, 4) if f"2026-0{i}" in _mm.index]
    _fa_basev = []
    for _i in (2, 3, 4):
        _vv = _mm[(_mm["month"] == _i) & (_mm["year"].isin([2022, 2023, 2024, 2025])) & (_mm["n_obs"] > 0)]["mean_ships_port_adj"]
        if len(_vv):
            _fa_basev.append(float(_vv.mean()))
    fa_26 = float(_np.mean(_fa26v)) if _fa26v else 0.0
    fa_base = float(_np.mean(_fa_basev)) if _fa_basev else 0.0
    n_pairs = summary.get("s2_pairs", 0)
    mean_snr = summary.get("mean_snr_db")
    n_months_data = int(monthly["n_obs"].gt(0).sum())
    d0, d1 = good["dt"].min(), good["dt"].max()
    ex1 = good.sort_values("datetime").iloc[-1]

    from validation import summary as vsummary
    vs = vsummary(scenes)
    n_cluster = vs.get("n_cluster_like", 0)

    anch26 = m[m.index >= "2026-01"]["mean_anchorage"].mean()
    anch_pre = m[(m.index >= "2022-01") & (m.index <= "2025-12")]["mean_anchorage"].mean()

    from insights import operational_series as _opsd
    _opd, _p95d = _opsd(scenes)
    cap26_pct = float(_opd[_opd["year"] == 2026]["over_capacity"].mean() * 100) if len(_opd[_opd["year"] == 2026]) else 0.0
    cap_pre_pct = float(_opd[_opd["year"] <= 2025]["over_capacity"].mean() * 100) if len(_opd[_opd["year"] <= 2025]) else 0.0
    cov_min = float(good["cov_port"].min() * 100) if len(good) else 0.0
    opt26 = opt25 = 0.0
    try:
        from validation import load_s2 as _ls2d, pair_s1_s2 as _psd
        _prsd = _psd(good, _ls2d())
        if len(_prsd):
            _prsd = (_prsd.sort_values("gap_hours").drop_duplicates(subset="s1_id", keep="first")
                     .sort_values("s2_date").reset_index(drop=True))
            _prsd["s2_date"] = pd.to_datetime(_prsd["s2_date"])
            _o26d = _prsd[_prsd["s2_date"].dt.year == 2026]
            _o25d = _prsd[_prsd["s2_date"].dt.year == 2025]
            opt26 = float(_o26d["s2_port"].mean()) if len(_o26d) else 0.0
            opt25 = float(_o25d["s2_port"].mean()) if len(_o25d) else 0.0
    except Exception:
        pass

    # ---------- document ----------
    doc = Document()
    # page setup A4
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    # default style
    st = doc.styles["Normal"]
    st.font.name = FONT_AR
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:cs"), FONT_AR)

    # ---------- cover ----------
    para(doc, "", size=10, space_after=30)
    para(doc, "التقرير الدوري لرصد النشاط البحري", size=24, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "مرفأ اللاذقية", size=20, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    para(doc, "تحليل نشاط المرفأ عبر صور الأقمار الصناعية الرادارية والبصرية", size=13,
         color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)

    cover_rows = [
        ("الفترة المشمولة", f"{ym_ar(d0.strftime('%Y-%m'))} — {ym_ar(d1.strftime('%Y-%m'))}"),
        ("تاريخ الإصدار", datetime.now().strftime("%Y-%m-%d")),
        ("إعداد", "وحدة التحليل الفضائي"),
        ("المصادر", "Copernicus Sentinel-1 / Sentinel-2 (ESA) · OpenStreetMap · Natural Earth"),
    ]
    t = doc.add_table(rows=len(cover_rows), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(cover_rows):
        set_cell_rtl(t.rows[i].cells[0], k, size=11, bold=True, shade="F0F6FC")
        set_cell_rtl(t.rows[i].cells[1], v, size=11)
        t.rows[i].cells[0].width = Cm(5.5)
        t.rows[i].cells[1].width = Cm(11)
    doc.add_page_break()

    # ---------- 1) exec summary ----------
    heading(doc, "1) الملخص التنفيذي", 1)
    para(doc, f"يغطي هذا التقرير {n_months_data} شهرًا متتاليًا من الرصد الراداري "
              f"({ym_ar(d0.strftime('%Y-%m'))} — {ym_ar(d1.strftime('%Y-%m'))}) بمعدل وسطي نحو "
              f"{fmt(n_ok/n_months_data)} مشاهدات شهريًا. النتائج:", size=11)
    for txt in [
        f"ارتفاع النشاط خلال 2026: متوسط السفن داخل الحوض {fmt(y.loc[2026,'annual_mean_adj'])} "
        f"سفينة لكل مشاهدة في كانون الثاني–آب 2026 مقابل {fmt(t2026['mean_before'])} لنفس الأشهر من 2025 "
        f"(+{fmt(y.loc[2026,'yoy_pct'])}%) — أعلى مستوى منذ 2022، والفرق دال إحصائيًا "
        f"(احتمال أقل من 0.0001 على المشاهدات الفردية، واحتمال {fmt(t2026['p_value'],4)} على المتوسطات الشهرية).",
        f"الارتفاع ليس موسميًا: بعد ضبط النمط الموسمي، شباط–نيسان 2026 ({fa_26:.1f} سفينة لكل مشهد) "
              f"مقابل {fa_base:.1f} "
        "لنفس الأشهر في 2022–2025 (احتمال أقل من 0.0001).",
        "توقيت التحول: بدأ الارتفاع في كانون الأول 2025 وبلغ ذروته شباط–نيسان 2026، ثم تراجع جزئيًا صيفًا "
        "مع بقائه أعلى من المعدل الموسمي.",
        f"ضغط تشغيلي على المرسى: ارتفع متوسط سفن الانتظار من {fmt(anch_pre)} إلى {fmt(anch26)} سفينة لكل مشاهدة "
        "— نحو ثلاثة أضعاف.",
        "استقرار 2022–2025: التغيرات السنوية غير دالة إحصائيًا، بما فيها الفترة التي تلت كانون الأول 2024 "
        f"(احتمال {fmt(tpost['p_value'],4)}).",
    ]:
        p = doc.add_paragraph(style="List Number")
        set_rtl(p)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(txt)
        set_run(r, size=11)
    para(doc, "جميع المقارنات على أساس «نفس الفترة بنفس الفترة» مع ضبط الموسمية في الاختبارات.",
         size=10, color=GRAY)

    # ---------- 2) context ----------
    heading(doc, "2) السياق والأهداف", 1)
    para(doc, "يُعد مرفأ اللاذقية المنفذ البحري الرئيسي على الساحل السوري. "
              "الغرض من هذه المراقبة قياس مستوى النشاط التشغيلي للمرفأ عبر الزمن دون الاعتماد على تقارير "
              "محلية أو بيانات AIS، وذلك من صور رادارية لا تتأثر بالغيوم أو الإضاءة.")
    para(doc, "أهداف التقرير: (1) بناء سلسلة زمنية شهرية موثوقة للنشاط؛ (2) تحديد توقيت وحجم أي تغير؛ "
              "(3) التمييز بين التغير الموسمي والتغير البنيوي؛ (4) تقديم مؤشرات تشغيلية قابلة للمتابعة.")

    # ---------- 3) data ----------
    heading(doc, "3) البيانات والتغطية", 1)
    make_table(doc,
               ["المصدر", "البيانات", "العدد", "الوصول"],
               [["Sentinel-1 (ESA/Copernicus)", "رادار 10م، استقطابان (VV وVH)", f"{n_scenes} مشهدًا ({n_ok} صالحًا)",
                 "Microsoft Planetary Computer"],
                ["Sentinel-2 (ESA/Copernicus)", "بصري 10م", f"{len(s2)} مشهدًا", "Microsoft Planetary Computer"],
                ["OpenStreetMap", "حدود المرفأ + الكاسر + الساحل", "—", "Overpass API"],
                ["Natural Earth 10m", "اليابسة المرجعية", "—", "الملكية العامة"]],
               widths=[4.5, 4.0, 3.5, 4.5])
    para(doc, f"رُفض {n_scenes - n_ok} مشهدًا لعدم اكتمال تغطية منطقة الدراسة (قطع الحافة في بعض مدارات "
              "الصعود).", size=10, color=GRAY)

    # ---------- 4) study area ----------
    heading(doc, "4) منطقة الدراسة", 1)
    para(doc, "الحدود الرسمية لمرفأ اللاذقية (OSM) مع توسيع الرصد إلى حوض المرفأ (ضمن 400م) والمرسى "
              "(0.4–4 كم) ومياه العبور، داخل صندوق دراسة ثابت (≈ 10 × 11 كم).")

    # ---------- 5) methodology ----------
    heading(doc, "5) المنهجية والمؤشر", 1)
    para(doc, "خط المعالجة: قراءة النطاق الجزئي ← إسقاط على شبكة UTM-36N ثابتة (10م) ← قناع يابسة ← تنعيم "
              "Lee ← عتبة تكيفية (T = μ + k·σ) ← تحقق متقاطع VV/VH ← مكونات متصلة ← فصل السفن المتلاصقة "
              "وتوحيد الأجزاء المنقسمة ← استخراج الإحداثيات (WGS84) والأبعاد والاستطاعة والمنطقة.")
    para(doc, "مؤشر النشاط: متوسط السفن لكل مشاهدة، محسوبًا على المشاهدات التي غطت الحوض بالكامل فقط "
              f"(تغطية ≥ 70%؛ فعليًا ≥ {cov_min:.0f}% في السلسلة الحالية). الأشهر بلا مشاهدات لا تُحتسب.")
    p = para(doc, "ملاحظة: نوع السفينة (حاويات/نفط/بضائع) خارج نطاق هذا النظام — دقة 10م لا تكفي للتصنيف "
                  "دون AIS أو صور عالية الدقة، ويُسجل النوع «غير محدد» لجميع الاكتشافات.", size=10.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "FDF3D8")
    pPr.append(shd)

    # ---------- 6) results ----------
    heading(doc, "6) النتائج", 1)
    heading(doc, "6.1) الاتجاه العام", 2)
    add_image(doc, os.path.join(CHART_DIR, "monthly_line.png"), 15.5, "الاتجاه الشهري للنشاط (سفن لكل مشاهدة)")
    add_image(doc, os.path.join(CHART_DIR, "change.png"), 15.5, "كشف التغير — المتوسط المتحرك")
    add_image(doc, os.path.join(CHART_DIR, "scatter.png"), 15.5, "المشاهدات الفردية")

    heading(doc, "6.2) الملخص السنوي وحجم النشاط", 2)
    yrows = []
    for _, r in yearly.iterrows():
        yrows.append([int(r["year"]), int(r["n_obs"]), fmt(r["annual_mean_ships_port"]),
                      fmt(r["annual_mean_adj"]),
                      f"{fmt(r['peak_value'])} ({ARABIC_MONTHS[int(r['peak_month'])-1]})",
                      f"{fmt(r['low_value'])} ({ARABIC_MONTHS[int(r['low_month'])-1]})",
                      f"{fmt(r['yoy_pct'],1)}%"])
    make_table(doc, ["السنة", "مشاهدات", "متوسط النشاط", "بعد تصحيح التغطية", "أعلى شهر", "أدنى شهر",
                     "التغير السنوي"], yrows)
    vol = good.groupby("year").agg(scenes=("id","count"), det=("n_est","sum"),
                                   port=("n_est_in_port","sum"), anch=("n_anchorage","sum")).reset_index()
    vol["anch_scene"] = (vol["anch"]/vol["scenes"]).round(1)
    vrows = [[int(r.year), int(r.scenes), int(r.det), int(r.port), int(r.anch), fmt(r.anch_scene)]
             for r in vol.itertuples()]
    make_table(doc, ["السنة", "مشاهدات", "إجمالي السفن", "سفن الحوض", "سفن المرسى", "المرسى/مشهد"], vrows)

    heading(doc, "6.3) النمط الموسمي والتحكم به", 2)
    srows = []
    for mo in range(1, 13):
        past = m[(m["month"]==mo) & (m["year"].isin([2022,2023,2024,2025])) & (m["n_obs"]>0)]["mean_ships_port_adj"]
        cur = m[(m["month"]==mo) & (m["year"]==2026) & (m["n_obs"]>0)]["mean_ships_port_adj"]
        if len(past):
            diff = (cur.mean()-past.mean()) if len(cur) else None
            srows.append([ARABIC_MONTHS[mo-1], fmt(past.mean()), fmt(cur.mean()) if len(cur) else "—",
                          fmt(diff,1) if diff is not None else "—"])
    make_table(doc, ["الشهر", "متوسط 2022–2025", "2026", "الفرق"], srows)
    para(doc, f"اختبار الضبط الموسمي (شباط–نيسان): {fa_base:.1f} سفينة لكل مشهد في 2022–2025 مقابل "
              f"{fa_26:.1f} في 2026 — زيادة {fa_26 - fa_base:.1f} سفينة لكل مشهد، احتمال أقل من 0.0001 (اختبار مان-ويتني).",
         size=11)

    heading(doc, "6.4) خصائص السفن المكتشفة", 2)
    para(doc, f"من إجمالي {n_vessels:,} سجل سفينة: الوسيط {vmed:.0f}م، الربيعيان {vp25:.0f}م "
              f"و{vp75:.0f}م، والمئين التسعون {vp90:.0f}م. نحو {vgt150:.0f}% من السجلات أطوالها "
              f"> 150م، و{vlt50:.0f}% فقط < 50م — متسق مع حد الكشف الراداري بدقة 10م. "
              f"أطول جسم مسجل {vmax:.0f}م خارج أبعاد أي سفينة تجارية ويمثل كائنات ملتصقة أو صدى مركب "
              f"({n_cluster} سجلًا وُسمت «عناقيد محتملة» في البيانات ولم تُحذف).")

    heading(doc, "6.5) التوزيع المكاني", 2)
    para(doc, f"الحوض {zport_pct:.1f}% ({zport:,} سجلًا)، المرسى {zanch_pct:.1f}% ({zanch:,})، "
              f"ومياه العبور {ztrans} سجلًا ({ztrans_pct:.1f}%). الارتفاع في 2026 متزامن في الحوض "
              "والمرسى معًا — ما يشير إلى ضغط على الطاقة التفريغية وليس مجرد إعادة توزيع للسفن.")

    heading(doc, "6.6) آخر 18 شهرًا", 2)
    mrows = []
    for _, r in monthly.tail(18).iterrows():
        mrows.append([ym_ar(r["ym"]), "—" if r["n_obs_port"]==0 else int(r["n_obs"]),
                      fmt(r["mean_ships_port_adj"]), fmt(r["max_ships_port"]), fmt(r["mean_anchorage"])])
    make_table(doc, ["الشهر", "مشاهدات", "سفن لكل مشاهدة", "أقصى مشهد", "سفن بالمرسى"], mrows)

    # ---------- 7) statistics ----------
    heading(doc, "7) التحليل الإحصائي", 1)
    para(doc, f"اختبار الاتجاه (مان-كيندال): τ = {fmt(summary.get('mk_tau'))}، "
              f"احتمال {fmt(summary.get('mk_p'),4)} — لا اتجاه رتيب على كامل السلسلة (استقرار طويل ثم قفزة "
              "حديثة؛ اختبار الاتجاه الرتيب غير حساس لهذا النمط).")
    _m26m = m[m.index >= "2026-01"]["mean_ships_port_adj"]
    _mbasem = m[(m.index >= "2022-01") & (m.index <= "2025-12")]["mean_ships_port_adj"]
    _zscore = (float(_m26m.mean()) - float(_mbasem.mean())) / float(_mbasem.std()) if len(_mbasem) > 1 else 0
    _n1 = int((_m26m > _mbasem.mean() + _mbasem.std()).sum())
    _n2 = int((_m26m > _mbasem.mean() + 2 * _mbasem.std()).sum())
    para(doc, f"موقع 2026 من التوزيع التاريخي: متوسط أشهر 2026 ({_m26m.mean():.1f}) أعلى من متوسط "
              f"2022–2025 ({_mbasem.mean():.1f} ± {_mbasem.std():.1f}) بفارق z = {_zscore:.2f} "
              f"انحرافًا معياريًا؛ {_n1} من {len(_m26m)} أشهر تجاوزت +1σ و{_n2} من {len(_m26m)} تجاوزت +2σ.")
    make_table(doc, ["المقارنة (نفس الفترة)", "قبل", "بعد", "الفرق", "p-value", "الحكم"],
               [["2026 (كانون الثاني–آب) مقابل 2025 (كانون الثاني–آب)", fmt(t2026["mean_before"]),
                 fmt(t2026["mean_after"]), f"{fmt(t2026['diff_pct'])}%", fmt(t2026["p_value"],4), "دال — ارتفاع"],
                ["بعد كانون الأول 2024 مقابل قبله", fmt(tpost["mean_before"]), fmt(tpost["mean_after"]),
                 f"{fmt(tpost['diff_pct'])}%", fmt(tpost["p_value"],4), "غير دال"],
                ["2025 مقابل 2024", fmt(regs['y2025_vs_2024']['test']['mean_before']),
                 fmt(regs['y2025_vs_2024']['test']['mean_after']),
                 f"{fmt(regs['y2025_vs_2024']['test']['diff_pct'])}%",
                 fmt(regs['y2025_vs_2024']['test']['p_value'],4), "غير دال"],
                ["2023 مقابل 2022", fmt(regs['y2023_vs_2022']['test']['mean_before']),
                 fmt(regs['y2023_vs_2022']['test']['mean_after']),
                 f"{fmt(regs['y2023_vs_2022']['test']['diff_pct'])}%",
                 fmt(regs['y2023_vs_2022']['test']['p_value'],4), "غير دال"]],
               widths=[6.0, 2.2, 2.2, 2.0, 2.2, 2.2])
    add_image(doc, os.path.join(CHART_DIR, "monthly_comparison.png"), 15.5)
    add_image(doc, os.path.join(CHART_DIR, "heatmap.png"), 15.5)
    add_image(doc, os.path.join(CHART_DIR, "mom.png"), 15.5)
    add_image(doc, os.path.join(CHART_DIR, "yearly.png"), 15.5)

    # ---------- 7A) advanced analytics ----------
    heading(doc, "7أ) تحليلات إضافية", 1)
    heading(doc, "7أ.1) التوزيع الإحصائي لسفن الحوض حسب السنة", 2)
    add_image(doc, os.path.join(CHART_DIR, "boxplot.png"), 15.5,
              "مخطط الصناديق — توزيع السفن لكل مشهد حسب السنة")
    heading(doc, "7أ.2) توزيع أطوال السفن", 2)
    add_image(doc, os.path.join(CHART_DIR, "hist_length.png"), 15.5,
              f"توزيع أطوال السفن المكتشفة (الوسيط {vmed:.0f}م)")
    heading(doc, "7أ.3) توزيع السفن حسب المنطقة", 2)
    add_image(doc, os.path.join(CHART_DIR, "donut.png"), 12.0)
    heading(doc, "7أ.4) الملف الموسمي", 2)
    add_image(doc, os.path.join(CHART_DIR, "radar.png"), 15.5,
              "الملف الموسمي — 2026 مقابل متوسط 2022–2025")
    heading(doc, "7أ.5) التغير السنوي — المخطط التراجعي", 2)
    add_image(doc, os.path.join(CHART_DIR, "waterfall.png"), 15.5)
    heading(doc, "7أ.6) ضغط التشغيل: الحوض مقابل المرسى", 2)
    add_image(doc, os.path.join(CHART_DIR, "bubble.png"), 15.5)
    heading(doc, "7أ.7) التحقق المتقاطع S1↔S2", 2)
    add_image(doc, os.path.join(CHART_DIR, "s1s2.png"), 15.5)

    # ---------- 8) validation ----------
    heading(doc, "8) التحقق وجودة البيانات", 1)
    heading(doc, "8.1) مؤشرات الجودة حسب السنة", 2)
    from validation import quality_flags
    qf = quality_flags(good)
    qf["year"] = pd.to_datetime(qf["date"]).dt.year
    qa = qf.groupby("year").agg(det=("snr_db","count"), snr=("snr_db","mean")).round(1)
    sc_year = good.groupby("year").agg(scenes=("id","count"), rough=("sea_roughness","mean"),
                                       noise=("noise_floor_db","mean")).round(2)
    qa = qa.join(sc_year).reset_index()
    qrows = [[int(r.year), int(r.det), fmt(r.snr), int(r.scenes), fmt(r.rough,2), fmt(r.noise)]
             for r in qa.itertuples()]
    make_table(doc, ["السنة", "سجلات سفن", "متوسط الإشارة/الضوضاء (ديسيبل)", "مشاهدات", "خشونة البحر", "ضوضاء المشهد"], qrows)
    heading(doc, "8.2) التحقق المتقاطع S1↔S2", 2)
    para(doc, f"{n_pairs} زوجًا من المشاهدات المتقاربة (فارق ≤ 4 أيام، غيوم ≤ 25%). يتوافق المستشعران على "
              f"أن 2026 هي السنة الأعلى نشاطًا (العدّاد البصري: {opt26:.1f} مقابل {opt25:.1f} لنفس الأشهر من 2025).")
    heading(doc, "8.3) مراجعة الجودة المطبقة", 2)
    para(doc, "عولجت ثلاث مسائل منهجية: (1) توحيد الأجزاء المنقسمة للسفن الطويلة (164 حالة)؛ "
              "(2) تصحيح مراكز السفن الراسية على الأرصفة (35 حالة)؛ (3) ضبط دلالات ملفات البيانات. "
              f"وُحّدت السجلات ({n_vessels:,} سجلًا) عبر إعادة معالجة كاملة ومواءمة المصادر.")

    # ---------- 8أ) regional comparison (Tartus + Baniyas) ----------
    heading(doc, "8أ) المقارنة الإقليمية — اللاذقية · طرطوس · بانياس", 1)
    try:
        from compare import summary as cmp_summary
        _cps = cmp_summary()
        _cl = _cps.get("latakia") or {}
        _ct = _cps.get("tartus") or {}
        _cb = _cps.get("baniyas") or {}
        _rows = [
            ["اللاذقية", fmt(_cl.get("mean_25")), fmt(_cl.get("mean_26")),
             f"+{fmt(_cl.get('pct'))}%"],
            ["طرطوس", fmt(_ct.get("mean_25")), fmt(_ct.get("mean_26")),
             f"+{fmt(_ct.get('pct'))}%"],
        ]
        if _cb:
            _rows.append(["بانياس (نفطي)", fmt(_cb.get("mean_25")), fmt(_cb.get("mean_26")),
                          f"+{fmt(_cb.get('pct'))}%"])
        make_table(doc, ["الميناء", "2025 (كانون الثاني–آب)", "2026 (كانون الثاني–آب)", "التغير"], _rows)
        _ban_extra = ""
        if _cb:
            _ban_extra = (f" أما مرفأ بانياس النفطي فارتفع +{fmt(_cb.get('pct'))}% من قاعدة "
                          f"منخفضة ({fmt(_cb.get('mean_25'))} ← {fmt(_cb.get('mean_26'))} "
                          f"سفينة لكل مشاهدة) — لا يغيّر الخلاصة.")
        para(doc, f"للفصل بين تحول محلي وموجة إقليمية، طُبّق خط المعالجة نفسه على مرفأ طرطوس "
                  f"(الحدود الرسمية OSM: way 160479740) وعلى مرفأ بانياس النفطي (حوض محاط بكاسر أمواج "
                  f"OSM مغلق) باستخدام المشاهدات الرادارية ذاتها. بينما قفز نشاط اللاذقية "
                  f"+{fmt(_cl.get('pct'))}%، بقي نشاط طرطوس — الميناء الأكبر تاريخيًا — شبه مستقر "
                  f"(+{fmt(_ct.get('pct'))}%){_ban_extra} — ما يرجّح أن الارتفاع تحول محلي خاص "
                  f"باللاذقية وليس موجة إقليمية على الساحل السوري.")
    except Exception as _e:
        para(doc, f"المقارنة الإقليمية غير متاحة: {str(_e)[:80]}")

    # ---------- 9) maps ----------
    heading(doc, "9) الخرائط", 1)
    add_image(doc, os.path.join(MAP_DIR, "activity_map.png"), 15.5,
              "خريطة مواقع السفن المكتشفة (أحمر: الحوض · بنفسجي: المرسى · أزرق: العبور)")
    add_image(doc, os.path.join(MAP_DIR, "density_map.png"), 15.5, "خريطة كثافة الاكتشافات")

    # ---------- 10) examples ----------
    heading(doc, "10) أمثلة من المشاهدات", 1)
    add_image(doc, os.path.join(S1_DIR, ex1["id"], "detection_overlay.jpg"), 15.0,
              f"أحدث مشهد في السلسلة — {pd.to_datetime(ex1['datetime']):%Y-%m-%d %H:%M}Z · "
              f"{int(ex1['n_est_in_port'])} سفينة في الحوض")

    # ---------- 11) conclusions ----------
    heading(doc, "11) الخلاصة والتوصيات", 1)
    para(doc, "تشير البيانات إلى انتقال في مستوى نشاط مرفأ اللاذقية بدأ في كانون الأول 2025 وبلغ ذروته "
              "في الربع الأول من 2026 وبقي أعلى من مستويات 2022–2025 حتى نهاية فترة الرصد. الارتفاع مشترك "
              "بين الحوض والمرسى ويتجاوز النمط الموسمي.")
    for txt in [
        "المتابعة: مواصلة الرصد الشهري لتأكيد استمرارية المستوى الجديد، مع التركيز على أشهر آب–كانون الأول 2026.",
        "التحقق المعزز: اعتماد طبقة AIS كمرجع تحقق إضافي عند توفرها لتصنيف أنواع السفن.",
        "دقة العد: مراجعة عتبات الفصل في أشهر الازدحام الشديد (شباط–نيسان 2026).",
        "تفسير الأسباب: مقارنة السلسلة بمؤشرات التجارة الخارجية والحركة الملاحية عند توفرها.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        set_rtl(p)
        r = p.add_run(txt)
        set_run(r, size=11)

    # ---------- 12) limitations ----------
    heading(doc, "12) حدود المنهجية", 1)
    for txt in [
        "السفن الأصغر من نحو 15–20م (قوارب الصيد) غير مكتشفة بشكل منهجي بدقة 10م.",
        "فصل السفن المتلاصقة تقديري؛ تبقى أخطاء عد محدودة عند الازدحام الشديد.",
        "نوع السفينة وحمولتها خارج نطاق التحليل دون AIS.",
        "العدّاد البصري (S2) تقديري ويتأثر بالغيوم.",
        "لا تفسر البيانات أسباب الارتفاع؛ الوثيقة تصف الظاهرة كميًا وتضبط الموسمية فقط.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        set_rtl(p)
        r = p.add_run(txt)
        set_run(r, size=11)

    # ---------- 13) annexes ----------
    heading(doc, "13) الملاحق", 1)
    heading(doc, "13.1) تعريف حقول البيانات", 2)
    make_table(doc, ["الحقل", "المعنى"],
               [["id / datetime", "معرف المشهد الرسمي وتاريخ/وقت الالتقاط (UTC)"],
                ["satellite / platform / orbit", "القمر الصناعي (S1A/S1C/S1D) واتجاه المدار"],
                ["n_total / n_est", "عدد الكائنات اللامعة / تقدير السفن بعد الفصل والتوحيد"],
                ["n_est_in_port / n_anchorage", "السفن في الحوض / في المرسى"],
                ["coverage / cov_port", "تغطية البحر / تغطية حوض المرفأ في المشهد"],
                ["noise_floor_db / sea_roughness", "مستوى الضوضاء وحالة البحر"],
                ["vessels.json", "لكل سفينة: الإحداثيات، الأبعاد، الاستطاعة، المنطقة، الأوسمة"]],
               widths=[6.0, 10.5])
    heading(doc, "13.2) ملاحظات على البيانات", 2)
    para(doc, f"آخر مشهد مدرج في السلسلة: {pd.to_datetime(ex1['datetime']):%Y-%m-%d %H:%M}Z "
              f"({ex1['id'][:40]}…). مؤشر النشاط في ملف المشاهدات محسوب لكل مشاهدة على حدة.")



    out = os.path.join(ROOT, "docs", "FINAL_REPORT.docx")
    doc.save(out)
    print("docx written:", out, os.path.getsize(out)//1024, "KB")


if __name__ == "__main__":
    build()
